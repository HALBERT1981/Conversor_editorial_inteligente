# -*- coding: utf-8 -*-
"""
parser.py
Lê o .docx e extrai uma estrutura intermediária (ParsedDocument) pronta para
virar EPUB: capítulos detectados por Heading 1, blocos de texto formatados,
imagens em ordem de aparição (para detecção automática de capa) e metadados
básicos (título/autor) lidos das propriedades do documento.

Limitações conhecidas do MVP: tabelas do Word não são convertidas (apenas
parágrafos e imagens inline); cabeçalhos/rodapés não entram no conteúdo.
"""
from dataclasses import dataclass, field
from typing import List, Optional, Dict

from docx import Document
from docx.oxml.ns import qn

HEADING1_WORDS = ["heading 1", "título 1", "titulo 1", "heading1"]
HEADING2_WORDS = ["heading 2", "título 2", "titulo 2", "heading2"]
HEADING3_WORDS = ["heading 3", "título 3", "titulo 3", "heading3"]


@dataclass
class ImageAsset:
    id: str
    filename: str
    data: bytes
    media_type: str


@dataclass
class Run:
    text: str
    bold: bool = False
    italic: bool = False
    underline: bool = False


@dataclass
class Block:
    kind: str  # "heading2" | "heading3" | "paragraph" | "image" | "list_item"
    runs: List[Run] = field(default_factory=list)
    image_id: Optional[str] = None
    list_type: Optional[str] = None  # "bullet" | "number"


@dataclass
class Chapter:
    title: str
    blocks: List[Block] = field(default_factory=list)


@dataclass
class ParsedDocument:
    title: str
    author: str
    chapters: List[Chapter]
    images: List[ImageAsset]
    cover_candidate: Optional[ImageAsset]
    heading1_found: bool
    style_warnings: List[str] = field(default_factory=list)


def _style_name(paragraph) -> str:
    try:
        return (paragraph.style.name or "").strip().lower()
    except Exception:
        return ""


def _list_type(paragraph) -> Optional[str]:
    """Heurística simples: parágrafo com numPr é lista; tenta diferenciar
    numerada de marcador pelo nome do estilo (nem sempre é possível sem
    inspecionar numbering.xml em profundidade — suficiente para o MVP)."""
    p_pr = paragraph._p.pPr
    if p_pr is not None and p_pr.numPr is not None:
        style = _style_name(paragraph)
        if "number" in style or "número" in style or "numerad" in style:
            return "number"
        return "bullet"
    return None


def _paragraph_images(paragraph, doc_part, registry: Dict, counter: List[int]) -> List[ImageAsset]:
    found = []
    for run in paragraph.runs:
        for blip in run._element.findall(".//" + qn("a:blip")):
            r_id = blip.get(qn("r:embed"))
            if not r_id:
                continue
            try:
                image_part = doc_part.related_parts[r_id]
            except KeyError:
                continue
            key = image_part.partname
            if key in registry:
                found.append(registry[key])
                continue
            counter[0] += 1
            ext = image_part.content_type.split("/")[-1].replace("jpeg", "jpg")
            asset = ImageAsset(
                id=f"img_{counter[0]:03d}",
                filename=f"img_{counter[0]:03d}.{ext}",
                data=image_part.blob,
                media_type=image_part.content_type,
            )
            registry[key] = asset
            found.append(asset)
    return found


def _runs_from_paragraph(paragraph) -> List[Run]:
    runs = []
    for r in paragraph.runs:
        if not r.text:
            continue
        runs.append(Run(text=r.text, bold=bool(r.bold), italic=bool(r.italic), underline=bool(r.underline)))
    return runs


def _looks_like_heading(paragraph) -> bool:
    text = (paragraph.text or "").strip()
    if not text:
        return False
    if len(text) > 80:
        return False
    if len(text.split()) > 10:
        return False
    if any(p in text for p in [".", "?", "!"]):
        return False
    if any(r.bold for r in paragraph.runs):
        return True
    return False


def _is_heading_style(style_name: str) -> bool:
    return (
        any(w in style_name for w in HEADING1_WORDS)
        or any(w in style_name for w in HEADING2_WORDS)
        or any(w in style_name for w in HEADING3_WORDS)
    )


def parse_docx(path: str) -> ParsedDocument:
    doc = Document(path)
    core = doc.core_properties
    doc_title = (core.title or "").strip()
    doc_author = (core.author or "").strip()

    registry: Dict = {}
    counter = [0]
    images_in_order: List[ImageAsset] = []

    chapters: List[Chapter] = []
    current: Optional[Chapter] = None
    heading1_found = False
    heading2_found = False
    heading3_found = False
    style_warnings: List[str] = []
    possible_heading_candidates: List[str] = []

    for paragraph in doc.paragraphs:
        style = _style_name(paragraph)
        text = (paragraph.text or "").strip()

        para_images = _paragraph_images(paragraph, doc.part, registry, counter)
        for img in para_images:
            if img not in images_in_order:
                images_in_order.append(img)

        if any(w in style for w in HEADING1_WORDS):
            heading1_found = True
            current = Chapter(title=text or f"Capítulo {len(chapters) + 1}")
            chapters.append(current)
            continue

        if current is None:
            current = Chapter(title=doc_title or "Capítulo 1")
            chapters.append(current)

        if para_images:
            for img in para_images:
                current.blocks.append(Block(kind="image", image_id=img.id))
            if not text:
                continue

        if any(w in style for w in HEADING2_WORDS):
            heading2_found = True
            current.blocks.append(Block(kind="heading2", runs=_runs_from_paragraph(paragraph)))
            continue
        if any(w in style for w in HEADING3_WORDS):
            heading3_found = True
            current.blocks.append(Block(kind="heading3", runs=_runs_from_paragraph(paragraph)))
            continue

        if _looks_like_heading(paragraph) and not _is_heading_style(style):
            possible_heading_candidates.append(text)

        lt = _list_type(paragraph)
        if lt:
            current.blocks.append(Block(kind="list_item", runs=_runs_from_paragraph(paragraph), list_type=lt))
            continue

        if text:
            current.blocks.append(Block(kind="paragraph", runs=_runs_from_paragraph(paragraph)))

    if not heading1_found and heading2_found:
        style_warnings.append(
            "O documento contém Heading 2/3, mas não foi detectado Heading 1. Isso pode indicar hierarquia de títulos incorreta."
        )

    if possible_heading_candidates:
        candidates_text = "; ".join(possible_heading_candidates[:3])
        style_warnings.append(
            "Foram detectados possíveis títulos sem estilo Heading aplicado: "
            f"{candidates_text}{'...' if len(possible_heading_candidates) > 3 else ''}"
        )

    if not heading1_found and not heading2_found and possible_heading_candidates:
        style_warnings.append(
            "Verifique se os títulos do documento usam estilos Heading 1/2/3 no Word."
        )

    if not heading1_found and not heading2_found and not possible_heading_candidates:
        style_warnings.append(
            "Nenhum título claro foi encontrado. Verifique se o documento possui estilos Heading aplicados."
        )

    cover_candidate = images_in_order[0] if images_in_order else None

    return ParsedDocument(
        title=doc_title,
        author=doc_author,
        chapters=chapters,
        images=images_in_order,
        cover_candidate=cover_candidate,
        heading1_found=heading1_found,
        style_warnings=style_warnings,
    )
